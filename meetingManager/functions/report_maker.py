# -*- coding: utf-8 -*-
import os
from os import path, mkdir
import docx
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import datetime

now = datetime.datetime.now()

from researchManager.models import Files
BD = Files.objects


class ReportInfo(object):
	def __init__(
		self, mail, organization, id_meeting, date, time,
		type_research, description, main_researcher, doc_list, 
		id_research_in_bd, expert, predsedatel, 
		secretar, year, podrazdelenie, phone_num,
	):
		"""Constructor"""
		self.mail = mail
		self.organization = organization
		self.id_meeting = id_meeting
		self.date = date
		self.time = time
		self.type_research = type_research
		self.description = description
		self.main_researcher = main_researcher
		self.podrazdelenie = podrazdelenie
		self.doc_list = doc_list
		self.id_research_in_bd = id_research_in_bd
		self.expert = expert
		self.predsedatel = predsedatel
		self.secretar = secretar
		self.year = year
		self.phone_num = phone_num

	def getDocListString(self):
		docListString = ''
		count = 1
		for doc in self.doc_list:
			docListString += str(count) + ') ' + doc + '\n\n'
			count += 1
		return docListString

	def getBodyRecords(self):
		records = [
		    ('Вопрос об одобрении ', 'инициативного исследования'),
		    ('Название исследования: ', self.description),
		    ('Главный исследователь:  ', self.main_researcher),
		    ('Представлены документы: ', self.getDocListString()),
		]
		if self.podrazdelenie:
			records.append(('Наименование подразделения: ', self.podrazdelenie))
		return records
	
	def getEndRecords(self):
		records = (
			('Заявитель:', str(self.main_researcher)),
			('Заявка:', '№' + str(self.id_research_in_bd) + ' от ' + str(self.date)),
			('Эксперт:', str(self.expert)),
		)
		return records

def createDoc(reportInfoList, id_meeting, date, time, randNumber):
	document = Document()
	document.add_heading('ПОВЕСТКА ОЧЕРЕДНОГО ЗАСЕДАНИЯ КОМИТЕТА ПО ЭТИКЕ при ФГБУ «НМИЦ онкологии им. Н.Н. Петрова» Минздрава России', 0)
	document.add_heading('', 0)
	records = (
		 ('Заседание №:', id_meeting),
	    ('Место проведения заседания:', 'ФГБУ «НМИЦ онкологии им. Н.Н. Петрова» Минздрава России'),
	    ('Дата:', date),
	    ('Время:', time),
	)
	table = document.add_table(rows=1, cols=2)
	for type, text in records:
	    row_cells = table.add_row().cells
	    row_cells[0].text = type
	    row_cells[1].text = text
	document.add_heading('', 0)
	document.add_page_break()
	for reportInfo in reportInfoList:
		document.add_paragraph(reportInfo.type_research)
		records = reportInfo.getBodyRecords()
		table = document.add_table(rows=1, cols=2)
		for type, text in records:
			row_cells = table.add_row().cells
			row_cells[0].text = type
			row_cells[1].text = text
		records = reportInfo.getEndRecords()
		table = document.add_table(rows=1, cols=2)
		for type, text in records:
			row_cells = table.add_row().cells
			row_cells[0].text = type
			row_cells[1].text = text
		document.add_heading('', 0)
		document.add_page_break()
		path = os.path.dirname(__file__) + '\\..\\..\\media\\reports\\'   + str(now.strftime("%d-%m-%Y-%H")) + '\\'
		pathFromBd = f'reports/{str(now.strftime("%d-%m-%Y-%H"))}/subpoena{randNumber}.docx'
		if not os.path.exists(path):
			os.mkdir(path)
	document.save(f'{path}/subpoena{randNumber}.docx')
	reportInfoList = []

def getReportInfoList(researchList, mail, organization, id_meeting, date, time): #На вход кидаем список исследований в заседании
	#Помещаем все в цикл по списку исследований в заседании
	#Убираем все от сюда до следущего коммента
	reportInfoList = []
	for research in researchList:
		type_research = research.getTypeRequest()
		description = research.getDescription()
		main_researcher = research.main_researcher
		podrazdelenie = research.division
		doc_list = []
		for file in BD.filter(research_id=research.id):
			filename = ''
			filename += file.name + ', '
			if file.date:
				filename += "дата: " + file.date + ", "
			if file.version:
				filename += "версия: " + file.version + ", "
			doc_list.append(filename)
		id_research_in_bd = research.id
		expert = 'None None.None.'
		predsedatel = 'None.None. None'
		secretar = 'Т.А. Кириллова'
		year = '2021'
		phone_num = '4399555 (3304)'
		# и заменяем на research.description...
		reportInfo = ReportInfo(
			mail, organization, id_meeting, date, time,
			type_research, description, main_researcher, doc_list, 
			id_research_in_bd, expert, predsedatel, 
			secretar, year, podrazdelenie, phone_num,
		)
		doc_list = []
		reportInfoList.append(reportInfo)	#Оставляем вне цикла для заполнения списка данными о всех исследованиях в заседании
	return reportInfoList


def createReport(researchList, meeting, randNumber):
	#Данные о заседании
	mail = 'lec@niioncologii.ru'
	organization = 'ФГБУ «НМИЦ онкологии им. Н.Н. Петрова» \nМинздрава России'
	id_meeting = str(meeting.id)
	date = str(meeting.date.day) + '.' + str(meeting.date.month) + '.' + str(meeting.date.year)
	time = str(meeting.time)
	#Конец данных
	reportInfoList = getReportInfoList(researchList, mail, organization, id_meeting, date, time)
	print(researchList[0].work_name)
	createDoc(reportInfoList, id_meeting, date, time, randNumber)